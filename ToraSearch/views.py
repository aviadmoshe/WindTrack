from django.shortcuts import render,redirect
from django.core.urlresolvers import reverse,reverse_lazy
from django.contrib.auth import authenticate, login, logout
from django.views import generic
from django.views.generic import View
from django.views.generic.edit import CreateView,UpdateView,DeleteView,FormView
from django.http import HttpResponseRedirect
from django.contrib.auth.models import User
from .models import Yeshiva, UpLoadToraText
from .forms import AddYeshivaForm, UserForm, UserLoginForm, UploadForm, UpdateToraForm
from .SearchFunctions import Search
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import Http404
from pagedown.widgets import PagedownWidget

class YeshivaDelete(DeleteView):
    template_name = 'torasearch/yeshiva_confirm_delete.html'

    def get_object(self, queryset=None):
        obj = super(YeshivaDelete, self).get_object()
        if not obj.owner == self.request.user:
            raise Http404
        account = Yeshiva.objects.get(id=obj.user.id)
        self.pk = account
        return obj

    def get_success_url(self):
        return reverse(
            'detail',
            kwargs={'pk': self.pk}
        )

class yeshiva_delete(DeleteView):
    model = Yeshiva
    success_url = reverse_lazy('index')

class yeshiva_update(UpdateView):
    model = Yeshiva
    success_url = reverse_lazy('index')

class IndexView(generic.ListView):
    #redirect_field_name = 'next'
    template_name = 'torasearch/index.html'

    def get_queryset(self):
        if self.request.user.is_authenticated():
            return Yeshiva.objects.filter(user = self.request.user)
        # else:
        #     return Yeshiva.objects.all()

class DetailView(generic.DetailView):
    model = Yeshiva
    template_name = 'torasearch/detail.html'

class yeshiva_create(LoginRequiredMixin,View):
    login_url = 'login'
    form_class = AddYeshivaForm
    template_name = 'torasearch/registration_form.html'

    def get(self,request):
        form = self.form_class(None)
        return render(request, self.template_name, {'form': form})

    def post(self, request):
        form = self.form_class(request.POST)

        if form.is_valid():
            temp = form.save(commit = False)
            temp.user = request.user
            temp.save()
            return redirect('index')

        return render(request, self.template_name, {'form': form})

'''
class yeshiva_create(CreateView):
    template_name = 'torasearch/yeshiva_form.html'
    model = Yeshiva
    fields = '__all__' #['rav_name','name']

'''

class UpLoadView(CreateView):
    model = UpLoadToraText
    template_name = 'torasearch/upload_form.html'
    fields = '__all__'


class RegisterView(View):
    form_class = UserForm
    template_name = 'torasearch/registration_form.html'

    def get(self,request):
        form = self.form_class(None)
        return render(request, self.template_name, {'form': form})

    def post(self, request):
        form = self.form_class(request.POST)

        if form.is_valid():
            user = form.save(commit = False)
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user.set_password(password)
            user.save()

            user = authenticate(username = username, password = password)
            if user is not None:
                if user.is_active:
                    login(request, user)
                    return redirect('index')

        return render(request, self.template_name, {'form': form})

class LoginView(View):
    title = 'login'
    form_class = UserLoginForm
    template_name = 'torasearch/registration_form.html'

    def get(self, request):
        form = self.form_class(None)
        return render(request, self.template_name, {'form': form, 'title': self.title})

    def post(self, request):
        form = self.form_class(request.POST)

        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']

            user = authenticate(username=username, password=password)
            if user is not None:
                if user.is_active:
                    login(request, user)
                    print('you are now logged in!')
                    return redirect('index')

        return render(request, self.template_name, {'form': form})

class LogOutView(View):
    def get(self, request):
        logout(request,)
        print('you are now logged out')
        return redirect('index')

class ReadTextView(View):
    form_class = UpdateToraForm
    model = UpLoadToraText
    template_name = 'torasearch/read_text_pagedown.html'
    def get(self, request, pk):
        file = self.model.objects.get(pk=pk).text
        print('file=',file)
        if str(file).split('.')[1] == 'docx':
            import docx
            filename = str(file)
            def getText(filename):
                doc = docx.Document(filename)
                fullText = []
                for para in doc.paragraphs:
                    fullText.append(para.text)
                return ''.join(fullText)
            #st = getText(file)
            st = 'test string'
        else:
            #http://stackoverflow.com/questions/13412924/how-to-get-txt-file-content-from-filefield
            file.open(mode = 'rb')
            st = file.readlines()# anither option:file.read().decode()
            #st = st.replace(r'\n', '<br>')
        return render(request, self.template_name, {'file_content': st})

class TextReadView(FormView):
    template_name = 'torasearch/read_text.html'
    form_class = UpdateToraForm
    model = UpLoadToraText
    def get(self, request, pk):
        file = self.model.objects.get(pk=pk).text
        print('file=',file)
        if str(file).split('.')[1] == 'docx':
            import docx
            filename = str(file)
            def getText(filename):
                doc = docx.Document(filename)
                fullText = []
                for para in doc.paragraphs:
                    fullText.append(para.text)
                return ''.join(fullText)
            #self.st = getText(file)
            self.st = 'test string'
        else:
            #http://stackoverflow.com/questions/13412924/how-to-get-txt-file-content-from-filefield
            file.open(mode = 'rb')
            self.st = file.readlines()# anither option:file.read().decode()
            #st = st.replace(r'\n', '<br>')
        return render(request, self.template_name, {'form': self.form_class})
    def get_context_data(self, **kwargs):
        context = super(TextReadView, self).get_context_data(**kwargs)
        context = {
            'file_content':self.st,

        }
        return context

class homepage(View):
    template_name = 'torasearch/homepage.html'
    def get(self,request):
        return render(request,self.template_name)

import json
from django.http import HttpResponse
from haystack.query import SearchQuerySet


def autocomplete(request):
    sqs = SearchQuerySet().autocomplete(content_auto=request.GET.get('q', ''))[:5]
    suggestions = [result.title for result in sqs]
    # Make sure you return a JSON object, not a bare list.
    # Otherwise, you could be vulnerable to an XSS attack.
    the_data = json.dumps({
        'results': suggestions
    })
    return HttpResponse(the_data, content_type='application/json')


def search(request):
    if request.method == 'GET':
        input = request.GET["keyword"]
        mak = Search(input)
    return render(request,'torasearch\search_answer.html',{'input' :len(mak)})
